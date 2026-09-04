"""文書生成・合成サービス（公開の汎用リファレンス実装 / 既定の合成バックエンド）。

Open GENAI の `procuretech-editor` から呼ばれる pluggable な生成/合成 API の
「そのまま動く」実装。LLM/Dify には依存せず、同梱の簡単なヒアリングシート
（`materials/hearing/hearing-sample.xlsx`）を読み取り、章別 Markdown を生成し、
Word(.docx) 合成まで一通り行える。テーマ固有の非公開サービス（例: 調達仕様書=spec-app）を
差し替える際の雛形であり、テーマ無しの「素の文書」の Word 化の既定バックエンドでもある。

契約:
- POST /generate            multipart: 任意キーの Excel / form: username, doc_type, options
                            -> {"request_id": "..."}
- GET  /status/{id}         -> {"status": processing|success|error, "progress": int}
- GET  /result/{id}         -> application/zip（section*.md, README.md, sections.json）
- POST /compose             JSON {"outputs":[{"name","sections":[{"filename","content"}]}],
                                   "assets": {"images/x.png": "<base64>"}}
                            -> application/zip（<name>.docx）。assets の画像は本文の
                               `![](相対パス)` に一致すれば .docx へ埋め込む。
- GET  /template/{key}      -> 同梱のヒアリングシート様式（xlsx）をダウンロード

`GENERATE_API_KEY` が設定されていれば `X-API-Key` を検証する。
ジョブ状態はメモリ保持（プロセス再起動で消える）。
"""

from __future__ import annotations

import base64
import io
import os
import re
import time
import uuid
import zipfile
from pathlib import Path
from typing import Any

from fastapi import FastAPI, Header, Request
from fastapi.responses import JSONResponse, Response

# API キー（旧名 GENERATE_SAMPLE_API_KEY も後方互換で参照）。
API_KEY = os.environ.get("GENERATE_API_KEY") or os.environ.get("GENERATE_SAMPLE_API_KEY", "")
# success になるまでの擬似処理時間（polling UI を確認できるように少し待たせる）。
PROCESS_SECONDS = float(
    os.environ.get("GENERATE_PROCESS_SECONDS")
    or os.environ.get("GENERATE_SAMPLE_PROCESS_SECONDS", "2")
)

# 同梱のヒアリングシート様式（key -> ファイル名）。/template/{key} で配信する。
HEARING_DIR = Path(os.environ.get("HEARING_DIR", "materials/hearing"))
TEMPLATES: dict[str, str] = {"hearing": "hearing-sample.xlsx"}

# ヒアリングシート（項目|値）の「項目」ラベル -> (section_key, 章タイトル, 出力ファイル名, 表示順)
FIELDS: list[tuple[str, str, str, str, int]] = [
    ("背景", "background", "背景", "section1.md", 2),
    ("目的", "purpose", "目的", "section2.md", 3),
    ("対象業務", "target", "対象業務", "section3.md", 4),
    ("主要要件", "requirements", "主要要件", "section4.md", 5),
    ("想定スケジュール", "schedule", "想定スケジュール", "section5.md", 6),
]
TITLE_LABEL = "案件名"

app = FastAPI(title="ProcureTech Generate", version="1.0.0")

# request_id -> {"created": float, "zip": bytes, "doc_type": str}
_JOBS: dict[str, dict[str, Any]] = {}


def _check_key(x_api_key: str | None) -> JSONResponse | None:
    if API_KEY and x_api_key != API_KEY:
        return JSONResponse(status_code=401, content={"error": "invalid api key"})
    return None


def _read_pairs(raw: bytes) -> dict[str, str]:
    """ヒアリングシート先頭シートの A 列(項目)/B 列(値) を dict にする。"""
    import openpyxl

    pairs: dict[str, str] = {}
    wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    try:
        ws = wb[wb.sheetnames[0]]
        for row in ws.iter_rows(values_only=True):
            if not row:
                continue
            label = str(row[0]).strip() if row[0] is not None else ""
            value = str(row[1]).strip() if len(row) > 1 and row[1] is not None else ""
            if label:
                pairs[label] = value
    finally:
        wb.close()
    return pairs


def _bullets_or_paragraph(value: str) -> str:
    """複数行の値は箇条書き、単一行はそのまま段落にする。"""
    lines = [ln.strip() for ln in value.splitlines() if ln.strip()]
    if len(lines) > 1:
        return "\n".join(f"- {ln}" for ln in lines)
    return value or "（未記入）"


def _build_zip(files: dict[str, bytes], doc_type: str) -> bytes:
    """アップロードされたヒアリングシートから章別 Markdown zip を作る。"""
    pairs: dict[str, str] = {}
    for data in files.values():
        try:
            pairs = _read_pairs(data)
        except Exception:  # noqa: BLE001
            pairs = {}
        break  # サンプルは先頭の 1 ファイルのみ使用

    title = pairs.get(TITLE_LABEL) or "サンプル調達案件"
    ts = time.strftime("%Y-%m-%d %H:%M:%S")

    outputs: dict[str, str] = {}
    manifest: list[dict[str, Any]] = [
        {"file": "README.md", "section_key": "readme", "title": "README", "order": 1}
    ]
    outputs["README.md"] = (
        f"# {title}\n\n"
        f"> このファイルは汎用生成サービス（procuretech-generate-app）が生成しました"
        f"（doc_type: `{doc_type}`）。\n"
        f"> 本番では非公開の生成サービスがテンプレート＋LLM/Dify で章別 Markdown を生成します。\n\n"
        f"- 生成日時: {ts}\n"
    )
    for label, key, section_title, filename, order in FIELDS:
        body = _bullets_or_paragraph(pairs.get(label, ""))
        outputs[filename] = f"# {section_title}\n\n{body}\n"
        manifest.append(
            {"file": filename, "section_key": key, "title": section_title, "order": order}
        )

    import json

    outputs["sections.json"] = json.dumps(
        {"theme": "sample", "sections": manifest}, ensure_ascii=False, indent=2
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, body in outputs.items():
            zf.writestr(name, body)
    return buf.getvalue()


# 画像のみの行（ブロック画像として大きく埋め込む）。
_IMAGE_LINE_RE = re.compile(r"^!\[[^\]]*\]\(\s*<?([^)>\s]+)>?(?:\s+\"[^\"]*\")?\s*\)$")
# 行内（インライン）画像。テキストと混在していても抽出できる。
_INLINE_IMAGE_RE = re.compile(r"!\[[^\]]*\]\(\s*<?([^)>\s]+)>?(?:\s+\"[^\"]*\")?\s*\)")


def _rel_of(match_group: str) -> str:
    return match_group.replace("\\", "/").lstrip("/")


def _add_code_block(doc: Any, lang: str, lines: list[str]) -> None:
    """フェンス付きコードブロックを等幅段落で出力する。

    Mermaid はサーバ側では描画できない（本来はクライアントが合成前に PNG 化して画像へ差し替える）。
    未変換のまま届いた場合の保険として、注記＋ソースを崩さず出力する。
    """
    from docx.shared import Pt

    if lang == "mermaid":
        note = doc.add_paragraph()
        run = note.add_run("【Mermaid 図（画像未変換のためソースを表示）】")
        run.italic = True
    para = doc.add_paragraph()
    run = para.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def _add_line_with_inline_images(doc: Any, line: str, assets: dict[str, bytes]) -> None:
    """テキストと行内画像が混在する行を、画像を埋め込みつつ 1 段落で出力する。"""
    from docx.shared import Cm

    matches = list(_INLINE_IMAGE_RE.finditer(line))
    if not matches:
        doc.add_paragraph(line)
        return
    para = doc.add_paragraph()
    last = 0
    for m in matches:
        pre = line[last : m.start()]
        if pre:
            para.add_run(pre)
        rel = _rel_of(m.group(1))
        data = assets.get(rel)
        if data:
            try:
                para.add_run().add_picture(io.BytesIO(data), width=Cm(12))
            except Exception:  # noqa: BLE001
                para.add_run(f"[画像: {rel}]")
        else:
            para.add_run(f"[画像: {rel}]")
        last = m.end()
    tail = line[last:]
    if tail:
        para.add_run(tail)


def _markdown_to_docx(
    name: str, sections: list[dict[str, Any]], assets: dict[str, bytes] | None = None
) -> bytes:
    """章（Markdown 文字列）を連結し、簡易パースで .docx を作る（python-docx）。

    spec-app（pandoc）と同等に、本文が参照する画像を assets（{相対パス: バイト列}）から
    埋め込む。画像のみの行はブロック画像、テキスト混在はインライン画像として配置する。
    Mermaid は合成前にクライアントが PNG 画像へ差し替える運用のため、ここでは通常画像として
    埋め込まれる（未変換で届いた場合はコードブロックとして安全に出力する）。
    """
    from docx import Document
    from docx.shared import Cm

    assets = assets or {}
    doc = Document()
    doc.add_heading(name, level=0)
    for sec in sections:
        content = str(sec.get("content") or "")
        in_code = False
        code_lang = ""
        code_lines: list[str] = []
        for raw_line in content.splitlines():
            stripped = raw_line.strip()
            if stripped.startswith("```"):
                if in_code:
                    _add_code_block(doc, code_lang, code_lines)
                    in_code, code_lang, code_lines = False, "", []
                else:
                    in_code, code_lang, code_lines = True, stripped[3:].strip().lower(), []
                continue
            if in_code:
                code_lines.append(raw_line)
                continue
            line = raw_line.rstrip()
            if not line.strip():
                continue
            m = _IMAGE_LINE_RE.match(line.strip())
            if m:
                rel = _rel_of(m.group(1))
                data = assets.get(rel)
                if data:
                    try:
                        doc.add_picture(io.BytesIO(data), width=Cm(15))
                        continue
                    except Exception:  # noqa: BLE001
                        pass
                doc.add_paragraph(f"[画像: {rel}]")
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.lstrip().startswith(("- ", "* ")):
                doc.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
            else:
                _add_line_with_inline_images(doc, line, assets)
        if in_code and code_lines:  # フェンス閉じ忘れの保険
            _add_code_block(doc, code_lang, code_lines)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _decode_assets(raw: Any) -> dict[str, bytes]:
    """compose の assets（{相対パス: base64}）を {相対パス: バイト列} へ変換する。"""
    out: dict[str, bytes] = {}
    if not isinstance(raw, dict):
        return out
    for rel, b64 in raw.items():
        if not isinstance(rel, str) or not isinstance(b64, str):
            continue
        try:
            out[rel.replace("\\", "/").lstrip("/")] = base64.b64decode(b64)
        except Exception:  # noqa: BLE001
            continue
    return out


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.get("/template/{key}")
def template(key: str, x_api_key: str | None = Header(default=None)) -> Response:
    err = _check_key(x_api_key)
    if err:
        return err
    filename = TEMPLATES.get(key)
    path = HEARING_DIR / filename if filename else None
    if not path or not path.is_file():
        return JSONResponse(status_code=404, content={"error": "template not found"})
    return Response(
        content=path.read_bytes(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/generate")
async def generate(
    request: Request, x_api_key: str | None = Header(default=None)
) -> JSONResponse:
    """テーマ非依存: multipart 内の任意のアップロードファイルを受け付ける。"""
    err = _check_key(x_api_key)
    if err:
        return err
    form = await request.form()
    files: dict[str, bytes] = {}
    for key, value in form.multi_items():
        if hasattr(value, "read"):  # UploadFile
            files[key] = await value.read()
    doc_type = str(form.get("doc_type") or "sample")
    if not files:
        return JSONResponse(status_code=400, content={"error": "入力ファイルがありません"})
    request_id = uuid.uuid4().hex
    _JOBS[request_id] = {
        "created": time.time(),
        "zip": _build_zip(files, doc_type),
        "doc_type": doc_type,
    }
    return JSONResponse(status_code=202, content={"request_id": request_id})


@app.get("/status/{request_id}")
def status(request_id: str, x_api_key: str | None = Header(default=None)) -> JSONResponse:
    err = _check_key(x_api_key)
    if err:
        return err
    job = _JOBS.get(request_id)
    if job is None:
        return JSONResponse(status_code=404, content={"error": "not found"})
    elapsed = time.time() - job["created"]
    if elapsed < PROCESS_SECONDS:
        pct = int(min(90, (elapsed / max(PROCESS_SECONDS, 0.001)) * 90))
        return JSONResponse(content={"status": "processing", "progress": pct})
    return JSONResponse(content={"status": "success", "progress": 100})


@app.get("/result/{request_id}")
def result(request_id: str, x_api_key: str | None = Header(default=None)) -> Response:
    err = _check_key(x_api_key)
    if err:
        return err
    job = _JOBS.get(request_id)
    if job is None:
        return JSONResponse(status_code=404, content={"error": "not found"})
    return Response(
        content=job["zip"],
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{request_id}.zip"'},
    )


@app.post("/compose")
async def compose(
    request: Request, x_api_key: str | None = Header(default=None)
) -> Response:
    """順序付き Markdown（出力ファイル毎）を .docx に合成して zip で返す。

    body.assets（{相対パス: base64}）に本文の `![](相対パス)` と一致する画像を渡すと、
    ブロック／インラインいずれの画像も .docx へ埋め込む（Mermaid はクライアントが PNG 化して
    画像として渡す運用）。
    """
    err = _check_key(x_api_key)
    if err:
        return err
    try:
        body = await request.json()
    except Exception:  # noqa: BLE001
        return JSONResponse(status_code=400, content={"error": "invalid json"})
    outputs = body.get("outputs") if isinstance(body, dict) else None
    if not isinstance(outputs, list) or not outputs:
        return JSONResponse(status_code=400, content={"error": "outputs がありません"})
    assets = _decode_assets(body.get("assets") if isinstance(body, dict) else None)

    buf = io.BytesIO()
    used: set[str] = set()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, o in enumerate(outputs, 1):
            if not isinstance(o, dict):
                continue
            name = str(o.get("name") or f"output{i}")
            sections = o.get("sections") or []
            if not isinstance(sections, list):
                sections = []
            docx = _markdown_to_docx(name, sections, assets)
            arc = f"{name}.docx"
            n = 2
            while arc in used:
                arc = f"{name}({n}).docx"
                n += 1
            used.add(arc)
            zf.writestr(arc, docx)
    return Response(
        content=buf.getvalue(),
        media_type="application/zip",
        headers={"Content-Disposition": 'attachment; filename="compose.zip"'},
    )
