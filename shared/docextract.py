"""添付ドキュメントのテキスト抽出（backend / rag-app 共通モジュール）。

ローカル LLM は PDF/Word/Excel 等を直接読めないため、テキストへ抽出して
プロンプトや RAG の知識ベースに流し込むための共通関数を提供する。

依存: pypdf / python-docx / openpyxl（各サービスの requirements に含める）。
"""

from __future__ import annotations

import base64
import io
import os
from typing import Any

# 抽出テキストの最大長（コンテキスト肥大を防ぐ）。ベクトル RAG / 簡易登録向け。
# 構造化索引では使わず、extract_doc_pages() を用いること。
MAX_DOC_CHARS = int(os.environ.get("MAX_DOC_CHARS", "30000"))

# チャット添付の全文抽出（その場マップリデュース）用のハード上限。
# 30k で黙って切らず、これを超えた場合のみ明示注記を付けて先頭を保持する。
MAX_CHAT_DOC_CHARS = int(os.environ.get("MAX_CHAT_DOC_CHARS", "500000"))

# 構造化取込のハード上限（黙って切らずエラーにする）
MAX_DOC_PAGES = int(os.environ.get("MAX_DOC_PAGES", "200"))
MAX_DOC_BYTES = int(os.environ.get("MAX_DOC_BYTES", str(20 * 1024 * 1024)))
# 非 PDF を合成ページに分割するときの目安文字数
SYNTHETIC_PAGE_CHARS = int(os.environ.get("SYNTHETIC_PAGE_CHARS", "3000"))


class DocExtractError(Exception):
    """構造化取込で上限超過など、黙って切り捨てできない失敗。"""

# UI の accept などに使える対応拡張子
SUPPORTED_DOC_EXTS = (
    ".pdf",
    ".docx",
    ".xlsx",
    ".txt",
    ".md",
    ".csv",
    ".tsv",
    ".html",
    ".htm",
    ".json",
    ".log",
)


def strip_base64_prefix(data: str) -> str:
    """`data:application/pdf;base64,xxxx` のような prefix を除去する。"""
    if data.startswith("data:"):
        comma = data.find(",")
        if comma != -1:
            return data[comma + 1 :]
    return data


def b64_to_bytes(data: str) -> bytes:
    return base64.b64decode(strip_base64_prefix(data))


def _docx_extract_text(raw: bytes) -> str:
    """docx から本文段落・表・ヘッダー/フッターを本文順に抽出する。

    python-docx の `document.paragraphs` は本文段落しか返さないため、表
    （テーブル）やテキストボックスに内容を持つ「様式」テンプレでは抽出漏れが
    起きる。ここでは本文ブロックを段落/表の順に走査し、表はセル単位（ネスト表も
    再帰）で拾う。ヘッダー/フッターも補足し、段落・表で何も取れないときのみ
    テキストボックス等を w:t 走査でフォールバック回収する。
    """
    import docx
    from docx.document import Document as _DocumentClass
    from docx.oxml.ns import qn
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    d = docx.Document(io.BytesIO(raw))

    def iter_block_items(parent: Any):
        if isinstance(parent, _DocumentClass):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            return
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def block_lines(container: Any) -> list[str]:
        lines: list[str] = []
        for block in iter_block_items(container):
            if isinstance(block, Paragraph):
                if block.text and block.text.strip():
                    lines.append(block.text)
            else:  # Table
                for row in block.rows:
                    cells: list[str] = []
                    prev: str | None = None
                    for cell in row.cells:
                        # 横結合セルは同一セルが繰り返されるため直前と同じ内容は畳む
                        val = " ".join(block_lines(cell)).strip()
                        if val and val != prev:
                            cells.append(val)
                        prev = val
                    line = "\t".join(cells)
                    if line.strip():
                        lines.append(line)
        return lines

    lines = block_lines(d)

    # ヘッダー / フッター（様式ではラベルが入ることがある）
    for section in d.sections:
        for hf in (section.header, section.footer):
            try:
                for p in hf.paragraphs:
                    if p.text and p.text.strip():
                        lines.append(p.text)
            except Exception:  # noqa: BLE001
                pass

    text = "\n".join(lines).strip()

    # 段落・表で何も取れない場合はテキストボックス等を w:t 走査でフォールバック
    if not text:
        parts = [t.text for t in d.element.iter(qn("w:t")) if t.text]
        text = "".join(parts).strip()
    return text


def _extract_raw_text(
    name: str, media_type: str, b64: str
) -> tuple[str | None, str | None]:
    """フォーマット判定して生テキストを抽出する（切り捨てなし）。

    戻り値 `(text, error)`:
    - 対応外（レガシー .doc/.xls 等）・base64 復号失敗 → `(None, None)`
    - 抽出中の例外 → `(None, "…失敗しました")`（呼び出し側でそのまま返せる文言）
    - 成功 → `(text, None)`
    """
    ext = name.lower().rsplit(".", 1)[-1] if "." in name else ""
    mt = media_type or ""
    try:
        raw = b64_to_bytes(b64)
    except Exception:  # noqa: BLE001
        return None, None

    try:
        if ext == "pdf" or mt == "application/pdf":
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw))
            text = "\n".join((p.extract_text() or "") for p in reader.pages)
        elif ext == "docx" or "wordprocessingml" in mt:
            text = _docx_extract_text(raw)
        elif ext == "xlsx" or "spreadsheetml" in mt:
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            lines: list[str] = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    lines.append("\t".join("" if c is None else str(c) for c in row))
            text = "\n".join(lines)
        elif ext in (
            "txt",
            "md",
            "csv",
            "tsv",
            "html",
            "htm",
            "json",
            "log",
        ) or mt.startswith("text/"):
            text = raw.decode("utf-8", "ignore")
        else:
            return None, None
    except Exception as e:  # noqa: BLE001
        return None, f"(添付ファイル {name} のテキスト抽出に失敗しました: {e})"

    return (text or ""), None


def extract_doc_text(name: str, media_type: str, b64: str) -> str | None:
    """添付ドキュメント(PDF/Word/Excel/テキスト)からテキストを抽出する。

    ベクトル RAG 簡易登録など従来経路向け。`MAX_DOC_CHARS` で切り捨てる。
    対応外（レガシー .doc/.xls 等）は None を返す。
    """
    text, error = _extract_raw_text(name, media_type, b64)
    if error is not None:
        return error
    if text is None:
        return None
    text = text.strip()
    if not text:
        return f"(添付ファイル {name} からテキストを抽出できませんでした)"
    if len(text) > MAX_DOC_CHARS:
        text = text[:MAX_DOC_CHARS] + "\n…(以下省略)"
    return text


def extract_doc_text_full(
    name: str, media_type: str, b64: str, *, max_chars: int | None = None
) -> str | None:
    """チャット添付向けの全文抽出（30k のサイレント切り捨てを行わない）。

    その場マップリデュースで扱うため、`MAX_DOC_CHARS` は適用しない。
    安全弁として `MAX_CHAT_DOC_CHARS`（既定 500,000）を超えた場合のみ、
    黙って捨てずに明示注記を付けて先頭を保持する。
    対応外（レガシー .doc/.xls 等）は None を返す。
    """
    text, error = _extract_raw_text(name, media_type, b64)
    if error is not None:
        return error
    if text is None:
        return None
    text = text.strip()
    if not text:
        return f"(添付ファイル {name} からテキストを抽出できませんでした)"
    cap = MAX_CHAT_DOC_CHARS if max_chars is None else max_chars
    if cap and cap > 0 and len(text) > cap:
        text = (
            text[:cap]
            + f"\n\n…（{name} は {cap} 文字を超えたため以降を省略しました。"
            "全文を対象にするにはナレッジ登録をご利用ください）"
        )
    return text


def _split_synthetic_pages(text: str) -> list[dict[str, Any]]:
    """非 PDF 向けに全文を保持したまま合成ページへ分割する。"""
    text = (text or "").strip()
    if not text:
        return []
    size = max(1, SYNTHETIC_PAGE_CHARS)
    pages: list[dict[str, Any]] = []
    start = 0
    page_no = 1
    while start < len(text):
        pages.append({"page": page_no, "text": text[start : start + size]})
        start += size
        page_no += 1
    return pages


def text_to_pages(text: str) -> list[dict[str, Any]]:
    """プレーンテキストを合成ページへ分割する（URL／簡易登録の全文保存用）。"""
    return _split_synthetic_pages(text)


def extract_doc_pages(name: str, media_type: str, b64: str) -> list[dict[str, Any]]:
    """構造化索引向けにページ単位で全文抽出する（MAX_DOC_CHARS は適用しない）。

    戻り値: [{"page": 1, "text": "..."}, ...]（page は 1 始まり）
    上限超過時は DocExtractError。対応外形式は空リスト。
    """
    ext = name.lower().rsplit(".", 1)[-1] if "." in name else ""
    mt = media_type or ""
    try:
        raw = b64_to_bytes(b64)
    except Exception as e:  # noqa: BLE001
        raise DocExtractError(f"base64 の復号に失敗しました: {e}") from e

    if len(raw) > MAX_DOC_BYTES:
        raise DocExtractError(
            f"ファイルサイズが上限（{MAX_DOC_BYTES} bytes）を超えています"
        )

    try:
        if ext == "pdf" or mt == "application/pdf":
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw))
            if len(reader.pages) > MAX_DOC_PAGES:
                raise DocExtractError(
                    f"ページ数が上限（{MAX_DOC_PAGES}）を超えています"
                    f"（{len(reader.pages)} ページ）"
                )
            pages: list[dict[str, Any]] = []
            for i, page in enumerate(reader.pages, start=1):
                pages.append({"page": i, "text": (page.extract_text() or "").strip()})
            if not any(p["text"] for p in pages):
                raise DocExtractError(
                    f"添付ファイル {name} からテキストを抽出できませんでした"
                )
            return pages

        # 非 PDF は MAX_DOC_CHARS を通さず全文抽出し、合成ページに分割する
        if ext == "docx" or "wordprocessingml" in mt:
            text = _docx_extract_text(raw).strip()
        elif ext == "xlsx" or "spreadsheetml" in mt:
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            lines: list[str] = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    lines.append("\t".join("" if c is None else str(c) for c in row))
            text = "\n".join(lines).strip()
        elif ext in (
            "txt",
            "md",
            "csv",
            "tsv",
            "html",
            "htm",
            "json",
            "log",
        ) or mt.startswith("text/"):
            text = raw.decode("utf-8", "ignore").strip()
        else:
            return []

        if not text:
            raise DocExtractError(
                f"添付ファイル {name} からテキストを抽出できませんでした"
            )
        pages = _split_synthetic_pages(text)
        if len(pages) > MAX_DOC_PAGES:
            raise DocExtractError(
                f"合成ページ数が上限（{MAX_DOC_PAGES}）を超えています"
                f"（{len(pages)} ページ相当）"
            )
        return pages
    except DocExtractError:
        raise
    except Exception as e:  # noqa: BLE001
        raise DocExtractError(
            f"添付ファイル {name} のテキスト抽出に失敗しました: {e}"
        ) from e
