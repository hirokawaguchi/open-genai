"""工程6: 文書読取（テキスト）と Vision 失敗時のフォールバック。"""

from __future__ import annotations

import asyncio
import base64
import sys
from pathlib import Path
from unittest.mock import AsyncMock, patch

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app import extract


def test_extract_text_document() -> None:
    payload = base64.b64encode("申請者 山田太郎".encode("utf-8")).decode("ascii")

    async def _run() -> None:
        result = await extract.extract_payload(
            kind="document",
            filename="note.txt",
            data=f"data:text/plain;base64,{payload}",
        )
        assert result["source"] == "docextract"
        assert "山田太郎" in result["extracted"]

    asyncio.run(_run())


def test_extract_rejects_unknown_kind() -> None:
    async def _run() -> None:
        try:
            await extract.extract_payload(kind="audio", filename="a.wav", data="AAAA")
        except ValueError as e:
            assert "kind" in str(e)
            return
        raise AssertionError("expected ValueError")

    asyncio.run(_run())


def test_extract_image_falls_back() -> None:
    payload = base64.b64encode(b"\x89PNG").decode("ascii")

    async def _run() -> None:
        with patch("app.extract.llm.chat_vision", new=AsyncMock(side_effect=RuntimeError("down"))):
            result = await extract.extract_payload(
                kind="image",
                filename="card.png",
                data=f"data:image/png;base64,{payload}",
            )
        assert result["source"] == "unavailable"
        assert result["extracted"] == ""
        assert "手入力" in (result.get("notes") or "")

    asyncio.run(_run())


def test_pages_to_guide_text_keeps_short_plain() -> None:
    assert extract.pages_to_guide_text([{"page": 1, "text": "転入の手引き"}]) == "転入の手引き"


def test_pages_to_guide_text_uses_headings() -> None:
    text = extract.pages_to_guide_text(
        [
            {
                "page": 1,
                "text": "# 手引き\n\n## 転入\n転入届を出します。\n\n## 転居\n転居届を出します。",
            }
        ]
    )
    assert "## 転入" in text
    assert "転入届を出します。" in text


def test_extract_docx() -> None:
    import io

    import docx

    doc = docx.Document()
    doc.add_heading("転入手引き", level=1)
    doc.add_paragraph("転入の場合は転入届が必要です。")
    buf = io.BytesIO()
    doc.save(buf)
    payload = base64.b64encode(buf.getvalue()).decode("ascii")

    async def _run() -> None:
        result = await extract.extract_payload(
            kind="document",
            filename="guide.docx",
            data=f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{payload}",
        )
        assert result["source"] == "docextract"
        assert "転入" in result["extracted"]

    asyncio.run(_run())


def test_extract_legacy_doc_unavailable() -> None:
    payload = base64.b64encode(b"\xd0\xcf\x11\xe0").decode("ascii")

    async def _run() -> None:
        result = await extract.extract_payload(
            kind="document",
            filename="old.doc",
            data=f"data:application/msword;base64,{payload}",
        )
        assert result["source"] == "unavailable"
        assert result["extracted"] == ""

    asyncio.run(_run())


if __name__ == "__main__":
    test_extract_text_document()
    test_extract_rejects_unknown_kind()
    test_extract_image_falls_back()
    test_pages_to_guide_text_keeps_short_plain()
    test_pages_to_guide_text_uses_headings()
    test_extract_docx()
    test_extract_legacy_doc_unavailable()
    print("ok")

